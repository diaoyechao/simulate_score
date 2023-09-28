import datetime
import json
import os
import subprocess

import docx
import pdfplumber
import requests

from rules import *


class BiddingDocumentParser:
    def __init__(self):
        self.bidding_announcement_parse_url = "http://192.168.60.139:8001/api/tender/tenderInformation"
        self.pattern_config = "config/pattern.txt"
        self.fixed_score_rule_patterns_config = "config/fixed_score_rule_patterns.json"

    @staticmethod
    def is_pdf(file_path):
        """
        判断招标文件是否是pdf文件
        :param file_path: 招标文件路径
        :return: True或者False
        """
        with open(file_path, 'rb') as file:
            header = file.read(4)
        return header == b'%PDF'

    @staticmethod
    def is_word(file_path):
        """
        判断招标文件是否是word文件
        :param file_path: 招标文件路径
        :return: True或者False
        """
        with open(file_path, 'rb') as file:
            header = file.read(2)
        return header == b'\xd0\xcf'

    @staticmethod
    def is_doc_or_docx(file_path):
        return file_path.split(".")[-1]

    @staticmethod
    def doc2docx(doc_path, docx_path):
        cmd = '/bin/libreoffice --headless --convert-to docx'.split() + [doc_path] + ['--outdir'] + [docx_path]
        p = subprocess.Popen(cmd, stderr=subprocess.PIPE, stdout=subprocess.PIPE)
        p.wait(timeout=30)

    @staticmethod
    def get_chapter_start_end_position(pdf):
        """
        获取招标文件评标办法章节的起始终止页
        :param pdf: 通过pdfplumber打开招标文件生成的pdf对象
        :return: 招标文件评标办法章节的起始终止页索引
        """
        start = 0
        end = 0
        chinese_characters_num_map = {"二": 2, "三": 3, "四": 4}
        num_chinese_characters_map = {value: key for key, value in chinese_characters_num_map.items()}
        start_pattern = "第.*章.*评标办法|第.*章.*评标及定标办法"
        end_pattern = ""
        for page_index in range(len(pdf.pages)):
            if page_index > 90:
                break
            if page_index < 5:
                continue
            page = pdf.pages[page_index]
            page_content = page.extract_text()
            page_content = "".join(
                ["".join(item.split()) for item in page_content.split("\n") if not re.search("非涉密|项目编号", item)])
            start_match = re.search(start_pattern, page_content[:60])
            if not end_pattern and start_match:
                start = page.page_number
                chinese_characters_pattern = r"\第(.+?)\章"
                chinese_characters_start_num = re.search(chinese_characters_pattern, start_match.group()).group(1)
                if chinese_characters_start_num.isdigit():
                    end_pattern = f"第{int(chinese_characters_start_num) + 1}章"
                else:
                    start_num = chinese_characters_num_map[chinese_characters_start_num]
                    end_num = start_num + 1
                    chinese_characters_end_num = num_chinese_characters_map[end_num]
                    end_pattern = f"第{chinese_characters_end_num}章"
            if end_pattern:
                if not page.extract_tables() and re.search(end_pattern, "".join(page_content[:60].split())):
                    end = page.page_number
                    break
        if start == 0:
            raise Exception("没有找到评标办法章节起始索引")
        if end == 0:
            raise Exception("没有找到评标办法章节终止索引")
        if start >= end:
            raise Exception("评标办法章节起始索引大于等于终止索引")
        return start, end

    @staticmethod
    def judge_evaluation_method_pdf(start, end, pdf):
        """
        从评标办法章节起始页到终止页中提取评估方法
        :param start: 起始页
        :param end: 终止页
        :param pdf: 通过pdfplumber打开招标文件生成的pdf对象
        :return: 评标办法
        """
        for j in range(start, end + 1):
            page = pdf.pages[j - 1]
            page_content = "".join(page.extract_text().split())
            if re.search("综合评估法|综合评分法|综合评价法", page_content):
                return re.search("综合评估法|综合评分法|综合评价法", page_content).group()
        return None

    @staticmethod
    def judge_evaluation_method_word(doc):
        paragraphs = doc.paragraphs
        for paragraph in paragraphs:
            paragraph_text = "".join(paragraph.text.split())
            if re.search("综合评估法|综合评分法|综合评价法", paragraph_text):
                return re.search("综合评估法|综合评分法|综合评价法", paragraph_text).group()
        return None

    @staticmethod
    def get_bid_rejection_rule_pdf(pdf):
        """
        获取废标规则
        :param pdf: 通过pdfplumber打开招标文件生成的pdf对象
        :return:
        """
        bid_rejection_rule = ""
        is_table = False
        for page_index in range(len(pdf.pages)):
            page = pdf.pages[page_index]
            tables = page.extract_tables()
            for table in tables:
                for table_row in table:
                    for item in table_row:
                        if item and re.search("废标", item):
                            bid_rejection_rule = "，".join(["".join(item.split()) for item in table_row if item])
                            is_table = True
                            break
            if not is_table:
                page_content = page.extract_text()
                if "废标" in page_content:
                    bid_rejection_rule_page_content = "".join(page_content.split())
                    # example0:[一二三四五六七八九十]、@\d+、@（\d+）
                    levels = ['[一二三四五六七八九十]、', '\\d+、']
                    result = []
                    level0_splits = re.split(levels[0], bid_rejection_rule_page_content)
                    for level0_split in level0_splits:
                        level1_split = re.split(levels[1], level0_split)
                        result.extend(level1_split)
                    if len(result) != 1:
                        start_index = 0
                        for index, item in enumerate(result):
                            if "废标" in item:
                                start_index = index
                        start_item = result[start_index]
                        end_index = start_index + 1
                        if end_index > len(result) - 1:
                            end_item = result[end_index]
                            bid_rejection_rule = bid_rejection_rule_page_content[bid_rejection_rule_page_content.find(
                                start_item):bid_rejection_rule_page_content.find(end_item)]
                        else:
                            bid_rejection_rule = bid_rejection_rule_page_content[
                                                 bid_rejection_rule_page_content.find(start_item):]
                        if bid_rejection_rule:
                            return bid_rejection_rule
                    # example1:\\d+.\\d+.\\d+
                    level_1 = "\\d+.\\d+.\\d+"
                    result = re.split(level_1, bid_rejection_rule_page_content)
                    if len(result) != 1:
                        bid_rejection_rule = "".join([item for item in result if re.search("废标", item)])
                        if bid_rejection_rule:
                            return bid_rejection_rule
        if is_table and not re.search("废标条款", bid_rejection_rule):
            split_items = re.split(r"\d、|\b\d{2}\b、", bid_rejection_rule)
            bid_rejection_rule = "".join([item for item in split_items if re.search("废标", item)])
        return bid_rejection_rule

    @staticmethod
    def get_bid_rejection_rule_word(doc):
        bid_rejection_rules = []
        paragraphs = doc.paragraphs
        all_content = ""
        for paragraph in paragraphs:
            paragraph_text = "".join(paragraph.text.split())
            all_content += paragraph_text
        all_contents = re.split("[一二三四五六七八]、|\\d+、|（\\d+）", all_content)
        for content in all_contents:
            if content and "废标" in content:
                bid_rejection_rules.append(content)
        return bid_rejection_rules

    @staticmethod
    def get_bidding_announcement_start_end_position(pdf):
        """
        获取招标文件招标公告章节的起始终止页
        :param pdf: 通过pdfplumber打开招标文件生成的pdf对象
        :return: 招标文件招标公告章节的起始终止页索引
        """
        start = 0
        end = 0
        chinese_characters_num_map = {"一": 1, "二": 2, "三": 3, "四": 4}
        num_chinese_characters_map = {value: key for key, value in chinese_characters_num_map.items()}
        start_pattern = "第.*章.*招标公告"
        end_pattern = ""
        for page_index in range(len(pdf.pages)):
            if page_index > 90:
                break
            page = pdf.pages[page_index]
            page_content = page.extract_text()
            if re.search(r"目录|目\s+录", page_content):
                continue
            start_match = re.search(start_pattern, "".join(page_content[:100].split()))
            if start_match:
                start = page.page_number
                chinese_characters_pattern = r"\第(.+?)\章"
                chinese_characters_start_num = re.search(chinese_characters_pattern, start_match.group()).group(1)
                if chinese_characters_start_num.isdigit():
                    end_pattern = f"第{int(chinese_characters_start_num) + 1}章"
                else:
                    start_num = chinese_characters_num_map[chinese_characters_start_num]
                    end_num = start_num + 1
                    chinese_characters_end_num = num_chinese_characters_map[end_num]
                    end_pattern = f"第{chinese_characters_end_num}章"
            if end_pattern:
                if re.search(end_pattern, "".join(page_content[:200].split())):
                    end = page.page_number
                    break
        if start == 0:
            raise Exception("没有找到招标公告章节起始索引")
        if end == 0:
            raise Exception("没有找到招标公告章节终止索引")
        if start > end:
            raise Exception("招标公告章节起始索引大于终止索引")
        return start, end

    def get_bidding_announcement_parse_result(self, bidding_announcement_content, title):
        """
        调用招标公告信息抽取接口获取解析结果
        :param bidding_announcement_content: 招标公告内容
        :param title: 标题
        :return: 解析结果
        """
        response = requests.post(url=self.bidding_announcement_parse_url,
                                 json={"content": bidding_announcement_content,
                                       "title": title})
        response_json = response.json()
        bidding_document_parse_result = {}
        for key, value in response_json["data"].items():
            if key == "5" or key == "6":
                continue
            if value["project_name"] == "行业分类":
                bidding_document_parse_result["industryType"] = value["label_name"]
            if value["project_name"] == "招标类型分类":
                bidding_document_parse_result["projectStage"] = value["label_name"]
            if value["project_name"] == "是否接受联合体投标":
                bidding_document_parse_result["isAccept"] = value["label_name"]
        proxyOrgTelephone = ""
        for json_item in response_json["data"]["5"]:
            if json_item["label_name"] == "项目名称":
                bidding_document_parse_result["projectName"] = json_item["value"]
            if json_item["label_name"] == "建设地点":
                bidding_document_parse_result["projectAddress"] = json_item["value"]
            if json_item["label_name"] == "项目招标编号":
                bidding_document_parse_result["projectNo"] = json_item["value"]
            if json_item["label_name"] == "招标单位联系电话":
                bidding_document_parse_result["projectTelephone"] = json_item["value"]
                bidding_document_parse_result["bidUnitTelephone"] = json_item["value"]
            if json_item["label_name"] == "招标单位":
                bidding_document_parse_result["bidUnit"] = json_item["value"]
            if json_item["label_name"] == "招标文件获取网址":
                bidding_document_parse_result["webSiteUrl"] = json_item["value"]
            if json_item["label_name"] == "招标单位地址":
                bidding_document_parse_result["bidUnitAddress"] = json_item["value"]
            if json_item["label_name"] == "招标代理机构":
                bidding_document_parse_result["proxyOrg"] = json_item["value"]
            if json_item["label_name"] == "招标代理机构地址":
                bidding_document_parse_result["proxyOrgAddress"] = json_item["value"]
            if json_item["label_name"] == "代理机构联系人":
                proxyOrgTelephone += json_item["value"]
            if json_item["label_name"] == "代理联系人电话":
                proxyOrgTelephone += json_item["value"]
            if json_item["label_name"] == "投标企业资质要求":
                bidding_document_parse_result["enterpriseQualification"] = json_item["value"]
            if json_item["label_name"] == "投标企业注册地要求":
                bidding_document_parse_result["enterpriseAddress"] = json_item["value"]
            if json_item["label_name"] == "投标企业备案要求":
                bidding_document_parse_result["enterpriseBackup"] = json_item["value"]
            if json_item["label_name"] == "投标企业财务要求":
                bidding_document_parse_result["enterpriseFinance"] = json_item["value"]
            if json_item["label_name"] == "投标企业业绩要求":
                bidding_document_parse_result["enterpriseKpi"] = json_item["value"]
            if json_item["label_name"] == "项目负责人资质要求":
                bidding_document_parse_result["projectLeaderQualification"] = json_item["value"]
            if json_item["label_name"] == "项目负责人业绩要求":
                bidding_document_parse_result["projectLeaderKpi"] = json_item["value"]
            if json_item["label_name"] == "项目负责人职称要求":
                bidding_document_parse_result["projectLeaderPm"] = json_item["value"]
            if json_item["label_name"] == "招标文件领取开始时间":
                if json_item.get("norm", None):
                    bidding_document_parse_result["startTime"] = json_item["norm"]
            if json_item["label_name"] == "招标文件领取结束时间":
                if json_item.get("norm", None):
                    bidding_document_parse_result["endTime"] = json_item["norm"]
        bidding_document_parse_result["proxyOrgTelephone"] = proxyOrgTelephone
        return bidding_document_parse_result

    def bidding_document_parse(self, file_path):
        bidding_document_parse_result = {
            "projectName": None,
            "projectNo": None,
            "projectTelephone": None,
            "bidControlPrice": None,  # todo
            "industryType": None,
            "projectStage": None,
            "bidUnit": None,
            "bidUnitAddress": None,
            "bidUnitTelephone": None,
            "proxyOrg": None,
            "proxyOrgAddress": None,
            "proxyOrgTelephone": None,
            "enterpriseQualification": None,
            "enterpriseAddress": None,
            "enterpriseBackup": None,
            "enterpriseFinance": None,
            "enterpriseKpi": None,
            "projectLeaderQualification": None,
            "projectLeaderKpi": None,
            "projectLeaderPm": None,
            "isAccept": None,
            "startTime": None,
            "contractLimit": None,
            "endTime": None,
        }
        if self.is_pdf(file_path) or os.path.splitext(file_path)[1] == ".pdf":
            with pdfplumber.open(file_path) as pdf:
                start, end = self.get_chapter_start_end_position(pdf)
                evaluation_method = self.judge_evaluation_method_pdf(start, end, pdf)
                bid_rejection_rule = self.get_bid_rejection_rule_pdf(pdf)
                bid_rejection_rules = [item for item in re.split("\\d+、", bid_rejection_rule)][1:]
                start_, end_ = self.get_bidding_announcement_start_end_position(pdf)
                bidding_announcement_content = ""
                title = pdf.pages[0].extract_text().split("\n")[0]
                for page_index in range(start_, end_):
                    page = pdf.pages[page_index - 1]
                    tables = page.find_tables()
                    if tables:
                        # todo
                        pass
                    page_content = page.extract_text()
                    bidding_announcement_content += page_content
                bidding_document_parse_result.update(
                    self.get_bidding_announcement_parse_result(bidding_announcement_content, title))
                bidding_document_parse_result["abandoneRuleList"] = bid_rejection_rules
                bidding_document_parse_result["scoreMethod"] = evaluation_method

        if self.is_word(file_path) or os.path.splitext(file_path)[1] in [".doc", ".docx"]:
            if self.is_doc_or_docx(file_path) == "doc":
                docx_path = file_path.split(".")[0] + ".docx"
                self.doc2docx(file_path, docx_path)
            else:
                docx_path = file_path
            doc = docx.Document(docx_path)
            evaluation_method = self.judge_evaluation_method_word(doc)
            bid_rejection_rule = self.get_bid_rejection_rule_word(doc)
            paragraphs = doc.paragraphs
            all_contents = "".join(["".join(paragraph.text.split()) for paragraph in paragraphs if paragraph])
            bid_index = all_contents.find("招标公告")
            notice_index = all_contents.find("投标须知")
            bidding_announcement_content = all_contents[bid_index:notice_index]
            title = ""  # todo
            bidding_document_parse_result.update(
                self.get_bidding_announcement_parse_result(bidding_announcement_content, title))
            bidding_document_parse_result["abandoneRuleList"] = bid_rejection_rule
            bidding_document_parse_result["scoreMethod"] = evaluation_method
        return bidding_document_parse_result

    def get_rating_sub_items_patterns(self):
        """
        读取评分子项识别pattern
        :return: 评分子项识别pattern
        """
        rating_sub_items_patterns = ""
        # 评分子项识别pattern
        pattern_file = open(self.pattern_config, encoding="utf-8")
        for pattern in pattern_file.readlines():
            pattern = pattern.strip()
            rating_sub_items_patterns += pattern
        pattern_file.close()
        return rating_sub_items_patterns

    def get_fixed_score_rule_patterns(self):
        fixed_score_rule_patterns = json.load(open(self.fixed_score_rule_patterns_config, encoding="utf-8"))
        return fixed_score_rule_patterns

    def recognize_rule(self, sub_item):
        """
        判断评分子项对应的规则
        :param sub_item: 评分子项
        :return: 如果有对应规则，则返回评分子项对应的规则，否则返回None。
        """
        sub_items = re.split("[，。；]", sub_item)
        a = [item for item in sub_items if re.match("(?=.*得\\d+分)(?!.*最多得\\d+分)(?!.*最多可得\\d+分)", item)]
        b = [item for item in sub_items if re.search("最多得\\d+分|最多可得\\d+分", item)]
        special_pattern = "满足|合理|可行|可靠|酌情赋分|最多得|优|良|得当"
        if len(re.findall("获得", sub_item)) >= 2 \
                or (len(a) >= 2 and not re.search(special_pattern, sub_item)) \
                or (len(a) >= 2 and len(b) >= 1):
            return "rule3", "满足【】条件得X分，满足【】条件得Y分，满足【】条件得Z分，..."

        for rule, patterns in self.get_fixed_score_rule_patterns().items():
            for pattern in patterns:
                if re.match(pattern, sub_item):
                    rule, rule_content = rule.split("-")
                    return rule, rule_content
        return None

    @staticmethod
    def split_items(item):
        if re.search("方式[一二三四五六七八九十]：", item):
            items = re.split("方式[一二三四五六七八九十]：", item)
        else:
            items = re.split("\\d+、|[一二三四五六七八九十]、|（\\d+）|[一二三四五六七八九十]标段业绩要求：|[\u2460\u2461\u2462\u2463\u2465]", item)
        if len(items) >= 2:
            return items
        return False

    def get_rating_sub_items(self, start, end, pdf):
        """
        提取出每一页的所有表格中的评分子项内容，添加到评分子项列表中
        :param start:评标办法章节起始位置
        :param end:评标办法章节终止位置
        :param pdf:pdf
        :return: 返回评分子项列表
        """
        rating_sub_items_patterns = self.get_rating_sub_items_patterns()
        rating_sub_items = []
        last_page_table_row = ""  # 被分割的上一页的表格行内容
        last_page_table_sign = False
        last_page_table_sign_index = None
        for j in range(start, end + 1):
            page = pdf.pages[j - 1]
            tables = page.extract_tables()
            # tables:每一页所有表格 [[[],[],...,[]],...,[[],[],...,[]]]
            # print(tables)
            # print("------")
            # 提取每一页每一个表格的评分子项内容
            for table_index, table in enumerate(tables):
                # table:每一页每一个表格 [[],[],...,[]]
                # print(table)
                # print("******")
                sign = False
                if re.search("评分分|分值|满分", "".join(str(table[0][-1]).split())) and len(
                        "".join(str(table[0][-1]).split())) < 5:
                    sign = True
                    if table_index == len(tables) - 1:
                        last_page_table_sign = True
                        last_page_table_sign_index = j
                    else:
                        last_page_table_sign = False
                if len([item for item in table[0] if
                        re.search("评审项|条款号|条款内容|编列内容|评分因素|评审内容|评审因素与评审标准", str(item))]) > 1:
                    table = table[1:]
                if last_page_table_sign and (
                        (None in table[0]) or ('' in table[0])) and j == last_page_table_sign_index + 1:
                    new_table = []
                    for table_row in table:
                        new_table.append([item for item in table_row if item and len(item) > 5])
                    table = new_table
                for table_row_index, table_row in enumerate(table):
                    # table_row:每一个表格的每一行
                    # print(table_row)
                    if table_row_index == len(table) - 1:
                        last_page_table_row = "".join(["".join(item.split()) for item in table_row if item])
                    # 处理表格中最后一列为评分分值的情况
                    if sign:
                        table_row = table_row[:-1]
                    if not table_row[-1]:
                        exist_rating_sub_item = ["".join(str(item).split()) for item in table_row if
                                                 re.search(rating_sub_items_patterns, "".join(str(item).split()))]
                        if exist_rating_sub_item:
                            rating_sub_items.extend(exist_rating_sub_item)
                        continue
                    # **********************
                    if not re.search(rating_sub_items_patterns, "".join(table_row[-1].split())):
                        if re.search("（\\d+分）", "".join(table_row[-2].split())):
                            filtered_table_row = ["".join(item.split()) for item in table_row[-2:]]
                            table_row_content = "，".join(filtered_table_row)
                            rating_sub_items.append(table_row_content)
                    # **********************
                    # 同一表格在不同页上被分割开的内容整合
                    # todo 还需优化
                    if table_row_index == 0:
                        if (None in table_row[:-1]) or ('' in table_row[:-1]) or last_page_table_sign:
                            last_page_table_row += "".join(table_row[-1].split())
                            # print(last_page_table_row)
                            search_rating_item = re.search(rating_sub_items_patterns, last_page_table_row)
                            # 判断是否存在评分子项，如果存在，则添加到评分子项列表中。
                            if search_rating_item and not re.search("投标报价", last_page_table_row):
                                rating_sub_items.append(last_page_table_row)
                            last_page_table_row = ""
                            continue
                    # 直接取每一个表格行的最后一个单元格的内容
                    table_row_content = "".join(table_row[-1].split())
                    search_rating_item = re.search(rating_sub_items_patterns, table_row_content)
                    # print(table_row_content)
                    # print("%%%%%%")
                    # 判断是否存在评分子项，如果存在，则添加到评分子项列表中。
                    if search_rating_item:
                        if re.search("投标报价", table_row_content):
                            continue
                        elif not self.recognize_rule(table_row_content) or re.search("有效期", table_row_content):
                            rating_sub_items.append(
                                "，".join(["".join(item.split()) for item in table_row if item and not item.isdigit()]))
                        else:
                            rating_sub_items.append(table_row_content)
        # print("")
        # for rating_sub_item in rating_sub_items:
        #     print(rating_sub_item)
        #     print("******")
        # exit()

        # 对小单元格继续进行拆分
        split_rating_sub_items = []
        for sub_item in rating_sub_items:
            count = 0
            if re.search("等次", sub_item):
                split_rating_sub_items.append(sub_item)
                continue
            items = self.split_items(sub_item)
            if items:
                for item in items:
                    if re.search(rating_sub_items_patterns, item):
                        count += 1
                        if count > 1:
                            break
                if count > 1:
                    for item in items:
                        if re.search(rating_sub_items_patterns, item):
                            split_rating_sub_items.append(item)
                else:
                    split_rating_sub_items.append(sub_item)
            else:
                split_rating_sub_items.append(sub_item)

        # 去重 去除多余信息：example：（注：...）
        filtered_split_rating_sub_items = []
        for index, item in enumerate(split_rating_sub_items):
            unuseful_mes_start_index = item.find("备注：") if item.find("备注：") != -1 else item.find("注：")
            if index == len(split_rating_sub_items) - 1:
                if unuseful_mes_start_index != -1:
                    filtered_split_rating_sub_items.append(item[:unuseful_mes_start_index])
                else:
                    filtered_split_rating_sub_items.append(item)
                break
            if item in split_rating_sub_items[index + 1]:
                continue
            else:
                if unuseful_mes_start_index != -1:
                    filtered_split_rating_sub_items.append(item[:unuseful_mes_start_index])
                else:
                    filtered_split_rating_sub_items.append(item)
        return list(dict.fromkeys(filtered_split_rating_sub_items))

    def score_analysis(self, file_path):
        final = {
            "success": True,
            "msg": "请求成功。",
            "code": 200,
            "data": {
                "fixedScoreResp": [],
                "subjectiveScoreResp": [],
                "businessScoreResp": {}
            },
            "timestamp": ""
        }
        if self.is_pdf(file_path) or os.path.splitext(file_path)[1] == ".pdf":
            with pdfplumber.open(file_path) as pdf:
                start, end = self.get_chapter_start_end_position(pdf)
                print(f"@1.评标办法章节开始位置{start}，结束位置{end}。")
                print("------")
                filtered_split_rating_sub_items = self.get_rating_sub_items(start, end, pdf)

            print("")
            print("@2.评分子项内容")
            for filtered_split_rating_sub_item in filtered_split_rating_sub_items:
                print(filtered_split_rating_sub_item)
                print("******")

            fixed_score_items = []
            subject_score_items = []
            # if os.path.exists("result.txt"):
            #     os.remove("result.txt")
            # w = open("./result.txt", "a", encoding="utf-8")
            for sub_item in filtered_split_rating_sub_items:
                # 可计算得分项
                # 判断评分子项对应的规则
                # w.write(sub_item + "\n")
                # print(sub_item)
                rule = self.recognize_rule(sub_item)
                # w.write(str(rule) + "\n")
                # w.write("\n")
                # print(rule)
                # print("&&&&&&")
                if rule:
                    if rule[0] == "rule1":
                        rule1_fixed_score_items = rule1(sub_item)
                        if rule1_fixed_score_items["scoreItem"]:
                            fixed_score_items.append(rule1_fixed_score_items)
                    if rule[0] == "rule2":
                        rule2_fixed_score_items = rule2(sub_item)
                        if rule2_fixed_score_items["scoreItem"]:
                            fixed_score_items.append(rule2_fixed_score_items)
                    elif rule[0] == "rule3":
                        rule3_fixed_score_items = rule3(sub_item)
                        if rule3_fixed_score_items["scoreItem"]:
                            fixed_score_items.append(rule3_fixed_score_items)
                else:
                    subject_score_items.append(sub_item)

            print("")
            print("@3.可计算得分项内容及其对应规则")
            fixedScoreResp = []
            for fixed_score_item in fixed_score_items:
                fixedScoreResp.append(fixed_score_item)
                print(fixed_score_item)
                print("******")
            final["data"]["fixedScoreResp"] = fixedScoreResp

            print("")
            print("@4.主观得分项及其对应最大分值")
            subjectiveScoreResp = []
            for subject_score_item in subject_score_items:
                subject_score_item_dict = {
                    "scoreItem": subject_score_item,
                    "maxScore": None,
                    "scoreType": 2,
                    "defaultScore": None
                }
                max_score_pattern = "最多得\\d+分|最高得(\\d+)分|满分\\d+分|最高不超过(\\d+)分|得\\d+分|得\\d+(\\.\\d+)?分" \
                                    "|得\\d+(\\.\\d+)?～\\d+分|得\\d+(\\.\\d+)?～\\d+(\\.\\d+)?分|为\\d+分|得\\d+-\\d+分|加\\d+分"
                max_score_iter = re.finditer(max_score_pattern, subject_score_item)
                scores = [max_score.group() for item in max_score_iter for max_score in
                          re.finditer("\\d+(\\.\\d+)?|\\d+", item.group())]

                scores = [int(score) if isinstance(is_integer_or_float(score), int) else float(score) for score in
                          scores]
                if scores:
                    subject_score_item_dict["maxScore"] = max(scores)
                subjectiveScoreResp.append(subject_score_item_dict)
            final["data"]["subjectiveScoreResp"] = subjectiveScoreResp

            businessScoreResp = {
                "businessScoreItem": None,
                "controlPrice": None,
                "upRuleRate": None,
                "downRuleRate": None,
                "basicRate": None,
                "basicPrice": None,
                "scoreType": 3
            }

            final["data"]["businessScoreResp"] = businessScoreResp

            # 获取当前时间
            current_time = datetime.datetime.now()
            # 将当前时间转换为时间戳（以秒为单位）
            timestamp = current_time.timestamp()
            final["timestamp"] = timestamp
            return final

        if self.is_word(file_path) or os.path.splitext(file_path)[1] in [".doc", ".docx"]:
            if self.is_doc_or_docx(file_path) == "doc":
                docx_path = file_path.split(".")[0] + ".docx"
                self.doc2docx(file_path, docx_path)
            else:
                docx_path = file_path
            doc = docx.Document(docx_path)
            return final


if __name__ == '__main__':
    file_path = "download_files/综合交通换乘中心全过程设计_1694688881460.pdf"
    biddingDocumentParser = BiddingDocumentParser()
    bidding_document_parse_result = biddingDocumentParser.bidding_document_parse(file_path)
    print(bidding_document_parse_result)
    score_analysis_result = biddingDocumentParser.score_analysis(file_path)
    print(score_analysis_result)
