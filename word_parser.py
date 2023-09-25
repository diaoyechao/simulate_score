import subprocess

import docx


# doc转docx
def doc2docx(docPath, docxPath):
    cmd = '/bin/libreoffice --headless --convert-to docx'.split() + [docPath] + ['--outdir'] + [docxPath]
    p = subprocess.Popen(cmd, stderr=subprocess.PIPE, stdout=subprocess.PIPE)
    p.wait(timeout=30)


# doc2docx("data/word/8.20招标文件-稽山街道2021年9月-2023年9月年度限额以下工程设计项目.doc", "data/word")


# doc = docx.Document("data/word/8.20招标文件-稽山街道2021年9月-2023年9月年度限额以下工程设计项目.docx")
# tables = doc.tables
# rating_sub_items = []
# for table in tables:
#     for row_index in range(len(table.rows)):
#         # 获取行对象
#         row = table.rows[row_index]
#         # 获取行对象文字信息
#         row_content = [cell.text for cell in row.cells]
#         # print(row_content)
#         for cell_content in row_content:
#             result = re.search(concat_patterns, cell_content)
#             if result:
#                 rating_sub_items.append("".join(cell_content.split()))
#
# for sub_item in rating_sub_items:
#     print(sub_item)
#     print("******")


# doc = docx.Document("data/word/8.20招标文件-稽山街道2021年9月-2023年9月年度限额以下工程设计项目.docx")
# paragraphs = doc.paragraphs
# all_contents = "".join(["".join(paragraph.text.split()) for paragraph in paragraphs if paragraph])
# bid_index = all_contents.find("招标公告")
# notice_index = all_contents.find("投标须知")
#
# bidding_announcement_content = all_contents[bid_index:notice_index]
# print(bidding_announcement_content)
